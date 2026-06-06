from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Helper: add a code-style paragraph
def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="007BFF" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5B)
    return h

# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SECURE USER REGISTRATION SYSTEM')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x2B, 0x5B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Server-Side Implementation Using Flask, SQLite & SHA-512')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Python (Flask) Edition — Network Security Assignment')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# Horizontal line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_paragraph()

# Course info table
info_data = [
    ['Course', 'Network Security'],
    ['Course Code', '115'],
    ['Submission Date', '11 June 2026'],
    ['Submitted To', 'Dr. Risala Tasin Khan, Professor'],
    ['', 'Jahangirnagar University'],
    ['Submitted By', 'ROBIUL HOSSAIN'],
    ['Student ID', '25306'],
    ['Program', 'PGDIT'],
    ['Institute', 'IIT, Jahangirnagar University'],
    ['Address', 'Savar, Dhaka 1342, Bangladesh'],
]
info_table = doc.add_table(rows=len(info_data), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_r = info_table.rows[i].cells[1]
    cell_l.text = ''
    cell_r.text = ''
    run_l = cell_l.paragraphs[0].add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(11)
    cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_r = cell_r.paragraphs[0].add_run(value)
    run_r.font.size = Pt(11)
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# Remove table borders for title info
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════
add_heading(doc, 'Table of Contents', level=1)
toc_items = [
    '1.  Abstract',
    '2.  Introduction',
    '3.  Project Objectives',
    '4.  Technology Stack',
    '5.  System Architecture & Design',
    '6.  Project Structure',
    '7.  Implementation Details',
    '    7.1  Database Schema',
    '    7.2  Password Hashing (SHA-512)',
    '    7.3  Input Validation',
    '    7.4  Session Management',
    '    7.5  Routing & Views',
    '8.  Security Analysis',
    '9.  Client-Side Features',
    '10. Testing & Results',
    '11. Comparison: Client-Side vs Server-Side',
    '12. Conclusion & Future Work',
    '13. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. ABSTRACT
# ══════════════════════════════════════════════════════════
add_heading(doc, '1. Abstract', level=1)
doc.add_paragraph(
    'This report presents the design and implementation of a Secure User Registration System '
    'built as a server-side web application using Python Flask, SQLite, and SHA-512 cryptographic '
    'hashing. The system provides a complete authentication workflow encompassing user registration, '
    'login, session management, and a protected dashboard. Security is addressed at multiple layers: '
    'passwords are hashed using SHA-512 before storage, all database queries use parameterized '
    'statements to prevent SQL injection, session cookies are cryptographically signed, and input '
    'validation is enforced on both the client side (JavaScript) and server side (Python). The '
    'application follows the Model-View-Controller (MVC) pattern using Flask\'s routing and Jinja2 '
    'templating. A dark mode toggle, real-time password strength meter, and inline field validation '
    'enhance the user experience. This project was developed as the final assignment for the '
    'Network Security course (Code: 115) at Jahangirnagar University.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ══════════════════════════════════════════════════════════
add_heading(doc, '2. Introduction', level=1)
doc.add_paragraph(
    'In the modern digital landscape, user registration and authentication systems form the '
    'gateway to nearly every web application. Ensuring the security of these systems is paramount, '
    'as they handle sensitive user credentials that, if compromised, can lead to identity theft, '
    'unauthorized access, and data breaches.'
)
doc.add_paragraph(
    'This project implements a secure user registration and login system using Python Flask on '
    'the server side. Unlike a purely client-side implementation where authentication logic and '
    'storage reside in the browser, this server-side approach ensures that sensitive operations '
    'such as password hashing and database access occur in a controlled, server-side environment '
    'inaccessible to end users.'
)
doc.add_paragraph(
    'The system demonstrates fundamental security principles including: cryptographic password '
    'storage using SHA-512 hashing, SQL injection prevention via parameterized queries, session-based '
    'authentication using signed cookies, and comprehensive input validation. Additionally, the '
    'application features a clean, responsive user interface with dark mode support, real-time '
    'password strength evaluation, and inline form validation for an improved user experience.'
)
doc.add_paragraph(
    'Built with Python 3.14+, Flask 3.x, SQLite3, and Jinja2 templating, the application follows '
    'industry best practices for web application security while remaining accessible as a learning '
    'tool for understanding authentication workflows.'
)

# ══════════════════════════════════════════════════════════
# 3. OBJECTIVES
# ══════════════════════════════════════════════════════════
add_heading(doc, '3. Project Objectives', level=1)
objectives = [
    ['Implement a secure user registration system with server-side password hashing using SHA-512.'],
    ['Create a login system with session-based authentication using Flask\'s signed cookies.'],
    ['Prevent SQL injection attacks through parameterized database queries.'],
    ['Enforce comprehensive input validation on both client and server sides.'],
    ['Provide a user-friendly interface with real-time password strength indicator and dark mode.'],
    ['Store user data persistently using SQLite with automatic database initialization.'],
    ['Demonstrate the security advantages of server-side authentication over client-side approaches.'],
    ['Implement duplicate email detection to prevent multiple registrations with the same email.'],
]
for obj in objectives:
    add_bullet(doc, obj[0])

# ══════════════════════════════════════════════════════════
# 4. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════
add_heading(doc, '4. Technology Stack', level=1)
tech_headers = ['Technology', 'Version', 'Purpose']
tech_rows = [
    ['Python', '3.14+', 'Runtime environment'],
    ['Flask', '3.x', 'Web framework (routing, sessions, templates)'],
    ['SQLite3', 'Built-in', 'Embedded database (file-based, no server needed)'],
    ['hashlib (SHA-512)', 'Built-in', 'Cryptographic password hashing'],
    ['Jinja2', 'Bundled with Flask', 'Server-side HTML templating'],
    ['HTML5 / CSS3', '—', 'Frontend structure and styling'],
    ['JavaScript (Vanilla)', '—', 'Client-side validation, dark mode, password strength'],
    ['python-docx', '1.2', 'Report generation (this document)'],
]
add_table(doc, tech_headers, tech_rows, col_widths=[4, 2.5, 8])

doc.add_paragraph()
doc.add_paragraph(
    'All dependencies are minimal. The only external package required is Flask, which is declared '
    'in requirements.txt. SQLite3 and hashlib are part of the Python standard library.'
)

# ══════════════════════════════════════════════════════════
# 5. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════
add_heading(doc, '5. System Architecture & Design', level=1)
doc.add_paragraph(
    'The application follows a classic client-server architecture using the Model-View-Controller '
    '(MVC) design pattern, adapted for Flask\'s request-response cycle.'
)

add_heading(doc, '5.1 Request-Response Flow', level=2)
flow_steps = [
    ['Browser requests a URL (e.g., /register).'],
    ['Flask matches the URL to a route handler in app.py.'],
    ['The route handler processes the request — validates input, interacts with the SQLite database, manages session state.'],
    ['Data is passed to a Jinja2 template for rendering.'],
    ['The template generates HTML, which is sent back to the browser.'],
    ['Browser renders the page; JavaScript handles client-side interactivity (validation, dark mode, password strength).'],
]
for step in flow_steps:
    add_bullet(doc, step[0])

add_heading(doc, '5.2 Architectural Layers', level=2)
arch_headers = ['Layer', 'Component', 'Responsibility']
arch_rows = [
    ['Presentation', 'Templates (HTML + Jinja2)', 'Rendering UI, displaying data, showing errors'],
    ['Client Logic', 'script.js', 'Form validation, dark mode toggle, password strength meter'],
    ['Controller', 'app.py (Routes)', 'Handling HTTP requests, coordinating logic'],
    ['Service', 'app.py (Validation + Hashing)', 'Input validation, SHA-512 hashing'],
    ['Data', 'SQLite (users.db)', 'Persistent storage of user records'],
]
add_table(doc, arch_headers, arch_rows, col_widths=[3, 4, 7.5])

# ══════════════════════════════════════════════════════════
# 6. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════
add_heading(doc, '6. Project Structure', level=1)
doc.add_paragraph('The project directory is organized as follows:')
code = '''secure_user_registration_system/
├── app.py                 Flask application (all routes, validation, DB logic)
├── requirements.txt       Python dependencies (Flask>=3.0)
├── users.db               SQLite database file (auto-created on first run)
├── README.md              Project documentation
├── templates/
│   ├── base.html          Base layout (dark mode toggle, footer, global JS)
│   ├── index.html         Home page with Register / Login buttons
│   ├── register.html      Registration form with all fields
│   ├── login.html         Login form with credential input
│   └── dashboard.html     Post-login dashboard showing all users
└── static/
    ├── style.css          Complete styling (light + dark mode)
    └── script.js          Client-side: validation, dark mode, password strength'''
add_code(doc, code)

doc.add_paragraph(
    'The templates directory contains five Jinja2 HTML files. base.html provides the common layout '
    '(dark mode toggle, footer with assignment details, and JavaScript includes), while the other '
    'templates extend it with page-specific content. The static directory holds CSS and JavaScript '
    'files served to the client.'
)

# ══════════════════════════════════════════════════════════
# 7. IMPLEMENTATION DETAILS
# ══════════════════════════════════════════════════════════
add_heading(doc, '7. Implementation Details', level=1)

# 7.1 Database Schema
add_heading(doc, '7.1 Database Schema', level=2)
doc.add_paragraph(
    'The database consists of a single table, users, with the following schema:'
)
add_code(doc, '''CREATE TABLE users (
    email         TEXT PRIMARY KEY,
    full_name     TEXT NOT NULL,
    phone         TEXT NOT NULL,
    password_hash TEXT NOT NULL
);''')
doc.add_paragraph(
    'The email field serves as the primary key, ensuring uniqueness. The password_hash field stores '
    'the SHA-512 digest of the user\'s password — never the plaintext password itself. The table is '
    'created automatically on the first database connection if it does not already exist (using '
    'CREATE TABLE IF NOT EXISTS).'
)

add_heading(doc, '7.1.1 Database Connection', level=3)
doc.add_paragraph(
    'The get_db() function establishes a connection to the SQLite database file and configures '
    'the row factory to return sqlite3.Row objects (dictionary-like access by column name):'
)
add_code(doc, '''def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("""CREATE TABLE IF NOT EXISTS users (
        email TEXT PRIMARY KEY,
        full_name TEXT NOT NULL,
        phone TEXT NOT NULL,
        password_hash TEXT NOT NULL
    )""")
    conn.commit()
    return conn''')

# 7.2 Password Hashing
add_heading(doc, '7.2 Password Hashing (SHA-512)', level=2)
doc.add_paragraph(
    'Password security is a critical aspect of any authentication system. This implementation uses '
    'SHA-512, a cryptographic hash function from the SHA-2 family, to hash passwords before storage. '
    'The hashing function is implemented as follows:'
)
add_code(doc, '''def hash_password(password):
    return hashlib.sha512(password.encode()).hexdigest()''')
doc.add_paragraph(
    'SHA-512 produces a fixed 512-bit (64-byte) hash value, typically represented as a 128-character '
    'hexadecimal string. This is a one-way function — given only the hash, it is computationally '
    'infeasible to recover the original password. When a user logs in, the system hashes the provided '
    'password and compares it with the stored hash. If they match, the password is correct.'
)
doc.add_paragraph(
    'Note: While SHA-512 is significantly more secure than MD5 or SHA-1, for production systems, '
    'purpose-built password hashing algorithms such as bcrypt, scrypt, or Argon2 are recommended '
    'as they include salting and are designed to be computationally expensive, making brute-force '
    'attacks more difficult.'
)

# 7.3 Input Validation
add_heading(doc, '7.3 Input Validation', level=2)
doc.add_paragraph(
    'The system implements dual-layer validation: client-side JavaScript provides immediate feedback '
    'to users, while server-side Python validation serves as the security boundary, ensuring that '
    'even if client-side validation is bypassed, invalid data is rejected.'
)

add_heading(doc, '7.3.1 Server-Side Validation', level=3)
doc.add_paragraph(
    'The validate() function in app.py performs comprehensive field-level validation using regular '
    'expressions and length checks:'
)
add_code(doc, '''def validate(field_id, value):
    msg = ""
    if field_id == "fullName":
        if not value: msg = "Full Name is required."
        elif not re.match(r"^[a-zA-Z\\s\\.]+$", value):
            msg = "Only letters, spaces, and dots allowed."
        elif len(value) < 2: msg = "Must be at least 2 characters."
    elif field_id in ("email", "loginEmail"):
        if not value: msg = "Email is required."
        elif not re.match(r"^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$", value):
            msg = "Enter a valid email address."
    elif field_id == "phone":
        if not value: msg = "Phone number is required."
        elif not re.match(r"^[\\d\\s\\+\\-\\(\\)]{7,15}$", value):
            msg = "Enter a valid phone number (7-15 digits)."
    elif field_id in ("password", "loginPassword"):
        if not value: msg = "Password is required."
        elif len(value) < 8: msg = "Must be at least 8 characters."
        elif not re.search(r"[A-Z]", value): msg = "Must contain an uppercase letter."
        elif not re.search(r"[a-z]", value): msg = "Must contain a lowercase letter."
        elif not re.search(r"[0-9]", value): msg = "Must contain a number."
        elif not re.search(r"[!@#$%^&*]", value):
            msg = "Must contain a special character (!@#$%^&*)."
    return msg''')

validation_rules_headers = ['Field', 'Rules']
validation_rules_rows = [
    ['Full Name', 'Required, letters/spaces/dots only, min 2 characters'],
    ['Email', 'Required, valid email format, must be unique'],
    ['Phone', 'Required, digits/spaces/+-/parentheses, 7-15 characters'],
    ['Password', 'Required, min 8 chars, must include uppercase, lowercase, digit, and special character (!@#$%^&*)'],
]
add_table(doc, validation_rules_headers, validation_rules_rows, col_widths=[3.5, 11])

doc.add_paragraph()
add_heading(doc, '7.3.2 Client-Side Validation', level=3)
doc.add_paragraph(
    'The script.js file mirrors the server-side validation rules for immediate user feedback. '
    'Fields are validated on input events, with error messages displayed inline and input borders '
    'changing color (red for invalid, green for valid). Form submission is prevented until all '
    'fields pass validation.'
)
add_code(doc, '''function validateField(id) {
    const input = document.getElementById(id);
    const error = document.getElementById(id + "Error");
    // ... validation rules mirroring server side ...
    error.innerText = msg;
    input.className = msg ? "invalid" : "valid";
    return msg === "";
}''')

# 7.4 Session Management
add_heading(doc, '7.4 Session Management', level=2)
doc.add_paragraph(
    'Flask\'s built-in session management is used to track authenticated users. When a user '
    'successfully logs in, their email is stored in the session object, which Flask serializes, '
    'signs with a random secret key, and stores in a client-side cookie:'
)
add_code(doc, '''# On login success:
session["email"] = email
return redirect(url_for("dashboard"))''')
doc.add_paragraph(
    'The secret key is randomly generated at application startup using os.urandom(24).hex(), '
    'producing a 48-character hexadecimal string. This key is used to cryptographically sign '
    'the session cookie, preventing tampering. Since the key changes on every server restart, '
    'existing sessions become invalid after the application restarts.'
)
doc.add_paragraph(
    'Protected routes like /dashboard check for the presence of session["email"]. If missing, '
    'the user is redirected to the login page. Logout simply removes the email from the session.'
)

# 7.5 Routing & Views
add_heading(doc, '7.5 Routing and Views', level=2)
routes_headers = ['Method', 'Path', 'Description']
routes_rows = [
    ['GET', '/', 'Home page with Register and Login buttons'],
    ['GET', '/register', 'Display registration form'],
    ['POST', '/register', 'Validate fields, hash password, insert into DB, redirect to login'],
    ['GET', '/login', 'Display login form'],
    ['POST', '/login', 'Validate credentials, set session, redirect to dashboard'],
    ['GET', '/dashboard', 'Display all registered users (requires authentication)'],
    ['GET', '/logout', 'Clear session and redirect to home'],
    ['Any', '/* (404)', 'Catch-all — redirect to home page'],
]
add_table(doc, routes_headers, routes_rows, col_widths=[2, 3, 9.5])

doc.add_paragraph()
doc.add_paragraph(
    'The dashboard route demonstrates session protection. It checks if the user is logged in '
    'before displaying sensitive data (all registered users and their password hashes). If a '
    'user attempts to access /dashboard without a valid session, they are redirected to /login.'
)

# ══════════════════════════════════════════════════════════
# 8. SECURITY ANALYSIS
# ══════════════════════════════════════════════════════════
add_heading(doc, '8. Security Analysis', level=1)

add_heading(doc, '8.1 Password Storage', level=2)
doc.add_paragraph(
    'Passwords are hashed using SHA-512 before being stored in the database. This ensures that '
    'even if the database is compromised, the original passwords cannot be easily recovered. '
    'The hash function is applied server-side, so the plaintext password never persists to disk.'
)

add_heading(doc, '8.2 SQL Injection Prevention', level=2)
doc.add_paragraph(
    'All database queries use parameterized statements with ? placeholders. This separates SQL '
    'code from user data, preventing malicious input from altering query structure. Example:'
)
add_code(doc, '''conn.execute(
    "SELECT password_hash FROM users WHERE email = ?", (email,)
)''')
doc.add_paragraph(
    'This approach ensures that user input is always treated as data, never as executable SQL code.'
)

add_heading(doc, '8.3 Session Security', level=2)
doc.add_paragraph(
    'Flask\'s session cookies are signed using a cryptographically random secret key generated '
    'at startup. This prevents attackers from forging or tampering with session data. The session '
    'cookie contains only the user\'s email address, and no sensitive information is stored client-side.'
)

add_heading(doc, '8.4 Validation as Security Boundary', level=2)
doc.add_paragraph(
    'Client-side validation is purely for user experience. The server independently validates '
    'all input, ensuring that even if an attacker bypasses the browser (e.g., using curl or '
    'Postman), all validation rules are enforced.'
)

add_heading(doc, '8.5 Known Limitations (Learning/Demo Project)', level=2)
limitations_headers = ['Limitation', 'Risk', 'Production Recommendation']
limitations_rows = [
    ['SHA-512 without salt', 'Vulnerable to rainbow table attacks if hash database is leaked', 'Use bcrypt, scrypt, or Argon2 with automatic salting'],
    ['No HTTPS', 'Credentials transmitted in plaintext over the network', 'Enable TLS/SSL certificate'],
    ['No rate limiting', 'Brute-force attacks on login endpoint', 'Implement rate limiting or account lockout'],
    ['No CSRF protection', 'Cross-Site Request Forgery on form submissions', 'Use Flask-WTF CSRF tokens'],
    ['Secret key per restart', 'All sessions invalidated on restart', 'Store persistent secret key in environment variable'],
    ['Single-factor auth', 'Only password-based authentication', 'Add 2FA (TOTP, SMS, etc.)'],
]
add_table(doc, limitations_headers, limitations_rows, col_widths=[4, 4.5, 6])

# ══════════════════════════════════════════════════════════
# 9. CLIENT-SIDE FEATURES
# ══════════════════════════════════════════════════════════
add_heading(doc, '9. Client-Side Features', level=1)

add_heading(doc, '9.1 Dark Mode Toggle', level=2)
doc.add_paragraph(
    'A fixed-position toggle button in the top-right corner switches between light and dark themes. '
    'The preference is persisted in localStorage, so the user\'s choice survives page reloads and '
    'browser sessions. CSS custom properties and body class toggling handle the visual transition.'
)

add_heading(doc, '9.2 Password Strength Meter', level=2)
doc.add_paragraph(
    'As the user types a password, a real-time indicator displays Weak, Medium, or Strong based on:'
)
strength_headers = ['Strength', 'Criteria']
strength_rows = [
    ['Weak', 'Password length < 6 characters'],
    ['Medium', 'Length >= 6 but missing one or more character categories'],
    ['Strong', 'Length >= 8, contains uppercase, lowercase, digit, and special character'],
]
add_table(doc, strength_headers, strength_rows, col_widths=[3, 11.5])

doc.add_paragraph()
doc.add_paragraph(
    'The strength is displayed in color-coded text (red, orange, green) below the password field.'
)

add_heading(doc, '9.3 Inline Field Validation', level=2)
doc.add_paragraph(
    'Each form field validates on input events. Invalid fields show a red border and an error message; '
    'valid fields show a green border. This provides immediate, intuitive feedback without requiring '
    'form submission.'
)

# ══════════════════════════════════════════════════════════
# 10. TESTING & RESULTS
# ══════════════════════════════════════════════════════════
add_heading(doc, '10. Testing & Results', level=1)
doc.add_paragraph(
    'The application was tested with the following scenarios to verify functionality and security:'
)

test_headers = ['Test Case', 'Input', 'Expected Result', 'Status']
test_rows = [
    ['Valid Registration', 'Name: John Doe, Email: john@test.com, Phone: +8801712345678, Password: Pass@123', 'User created, redirected to login', 'Pass'],
    ['Duplicate Email', 'Same email as above', 'Error: "Email already registered"', 'Pass'],
    ['Invalid Name', 'Name: John123', 'Error: "Only letters, spaces, and dots allowed"', 'Pass'],
    ['Invalid Email', 'Email: notanemail', 'Error: "Enter a valid email address"', 'Pass'],
    ['Weak Password', 'Password: short', 'Error: "Must be at least 8 characters"', 'Pass'],
    ['No Special Char', 'Password: Pass12345', 'Error: "Must contain a special character"', 'Pass'],
    ['Valid Login', 'Correct email and password', 'Session created, redirected to dashboard', 'Pass'],
    ['Invalid Login', 'Wrong password', 'Error: "Invalid Email or Password"', 'Pass'],
    ['SQL Injection Attempt', 'Email: admin\' OR \'1\'=\'1', 'No rows returned (parameterized query)', 'Pass'],
    ['Direct Dashboard Access', 'Access /dashboard without login', 'Redirected to /login', 'Pass'],
]
add_table(doc, test_headers, test_rows, col_widths=[4, 5.5, 4, 1.5])

# ══════════════════════════════════════════════════════════
# 11. COMPARISON
# ══════════════════════════════════════════════════════════
add_heading(doc, '11. Comparison: Client-Side vs Server-Side', level=1)
doc.add_paragraph(
    'This Python Flask implementation is the server-side successor to an earlier client-side '
    'JavaScript version. The following table compares the two approaches:'
)

comp_headers = ['Aspect', 'Original (JS / Client-Side)', 'Python (Flask / Server-Side)']
comp_rows = [
    ['Database', 'sql.js (WASM in browser) + IndexedDB', 'Python sqlite3 module + users.db file'],
    ['Hashing', 'Web Crypto API (crypto.subtle.digest)', 'hashlib.sha512() server-side'],
    ['Authentication', 'Query string in URL (insecure)', 'Flask signed session cookie'],
    ['Templates', 'Plain HTML files', 'Jinja2 templates (server-rendered)'],
    ['Validation', 'JavaScript only (bypassable)', 'Python server-side + JS client-side (dual-layer)'],
    ['Persistence', 'IndexedDB (browser storage, per-device)', 'On-disk .db file (shared across devices)'],
    ['Security', 'All logic exposed to user', 'Sensitive logic protected on server'],
]
add_table(doc, comp_headers, comp_rows, col_widths=[3, 5.5, 6])

doc.add_paragraph()
doc.add_paragraph(
    'The server-side approach provides significantly better security because sensitive operations '
    '(password hashing, database access) occur in a controlled environment. The client-side version, '
    'while useful for learning, exposes all authentication logic to the user and cannot truly secure '
    'any operation that must remain confidential.'
)

# ══════════════════════════════════════════════════════════
# 12. CONCLUSION & FUTURE WORK
# ══════════════════════════════════════════════════════════
add_heading(doc, '12. Conclusion & Future Work', level=1)

add_heading(doc, '12.1 Conclusion', level=2)
doc.add_paragraph(
    'This project successfully implements a secure user registration and authentication system '
    'using Python Flask with server-side SHA-512 password hashing. The system demonstrates '
    'several key security practices: parameterized SQL queries to prevent injection attacks, '
    'cryptographic password hashing, session-based authentication with signed cookies, and '
    'comprehensive dual-layer input validation.'
)
doc.add_paragraph(
    'The application provides a complete user workflow from registration through login to a '
    'protected dashboard, with a polished user interface featuring dark mode, real-time password '
    'strength evaluation, and inline form validation. The project serves as an effective '
    'demonstration of server-side security principles in web application development.'
)

add_heading(doc, '12.2 Future Work', level=2)
future_items = [
    ['Upgrade to bcrypt/Argon2 for password hashing with automatic salting.'],
    ['Add HTTPS support with TLS certificates.'],
    ['Implement rate limiting on login endpoints to prevent brute-force attacks.'],
    ['Add CSRF protection using Flask-WTF.'],
    ['Store the Flask secret key in an environment variable rather than generating per-restart.'],
    ['Implement email verification during registration.'],
    ['Add password reset functionality.'],
    ['Integrate Two-Factor Authentication (2FA) via TOTP.'],
    ['Add account lockout after multiple failed login attempts.'],
    ['Deploy using a production WSGI server (Gunicorn, Waitress) instead of Flask\'s dev server.'],
]
for item in future_items:
    add_bullet(doc, item)

# ══════════════════════════════════════════════════════════
# 13. REFERENCES
# ══════════════════════════════════════════════════════════
add_heading(doc, '13. References', level=1)
references = [
    'Flask Documentation — https://flask.palletsprojects.com/en/stable/',
    'Python hashlib Documentation — https://docs.python.org/3/library/hashlib.html',
    'Python sqlite3 Documentation — https://docs.python.org/3/library/sqlite3.html',
    'OWASP Password Storage Cheat Sheet — https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html',
    'OWASP SQL Injection Prevention — https://cheatsheetseries.owasp.org/cheatsheets/SQL_Injection_Prevention_Cheat_Sheet.html',
    'Jinja2 Template Documentation — https://jinja.palletsprojects.com/',
    'SHA-2 Cryptography — National Institute of Standards and Technology (NIST), FIPS PUB 180-4',
    'Web Crypto API — MDN Web Docs: https://developer.mozilla.org/en-US/docs/Web/API/Web_Crypto_API',
    'Python-Docx Library — https://python-docx.readthedocs.io/',
]
for i, ref in enumerate(references, 1):
    p = doc.add_paragraph(f'[{i}]  {ref}')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(2)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Secure_User_Registration_System_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
